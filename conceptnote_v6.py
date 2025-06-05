'''
pip install streamlit PyPDF2 pandas altair sentence-transformers nltk scikit-learn python-docx openpyxl pytesseract pdf2image Pillow
'''

import streamlit as st
import PyPDF2 # For text-based PDFs
import re 
import pandas as pd
import altair as alt
from typing import List, Dict, Tuple, Any
from io import BytesIO

# --- Libraries for new file formats and OCR ---
import openpyxl # For .xlsx files
from docx import Document as DocxDocument # For .docx files
from docx.table import _Cell # To check instance of cell
from docx.text.paragraph import Paragraph # To check instance of paragraph

from pdf2image import convert_from_bytes # For converting PDF pages to images for OCR
import pytesseract # Python wrapper for Tesseract OCR
from PIL import Image # For image manipulation with Pytesseract

# --- NLP and Similarity Libraries ---
from sentence_transformers import SentenceTransformer
import nltk
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import os # For Tesseract path if needed

# --- Configure Tesseract Path (IMPORTANT FOR SOME SYSTEMS, ESPECIALLY WINDOWS) ---
# If Tesseract is not in your system's PATH, uncomment and set the path below.
# Example for Windows:
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# Example for Linux (if not in default PATH):
# pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'

# --- Move st.set_page_config() here ---
st.set_page_config(layout="wide") # MUST BE THE FIRST STREAMLIT COMMAND

# --- Download NLTK resources (if not already downloaded) ---
@st.cache_resource
def download_nltk_resources():
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt', quiet=True)
    try:
        nltk.data.find('corpora/wordnet')
    except LookupError:
        nltk.download('wordnet', quiet=True)
    try:
        nltk.data.find('corpora/omw-1.4')
    except LookupError:
        nltk.download('omw-1.4', quiet=True)

download_nltk_resources()

# --- Load Sentence Transformer Model (this will download the model on first run) ---
@st.cache_resource
def load_similarity_model():
    try:
        model = SentenceTransformer('all-MiniLM-L6-v2')
        return model
    except Exception as e:
        st.error(f"Error loading SentenceTransformer model: {e}. Ensure you have an internet connection for the first run.")
        st.stop()

similarity_model = load_similarity_model()

# --- Similarity Threshold ---
SIMILARITY_THRESHOLD = 0.60 # 60% similarity

# --- Initial Keyword Matrix (Hardcoded from Concept Note) ---
# (Keyword matrix remains the same as in your provided code)
initial_indicator_keywords = {
    "Building Energy Efficiency": ["HVAC efficiency", "U-value", "building envelope", "building envelope efficiency", "energy demand", "energy performance", "energy simulation", "natural ventilation", "passive solar", "thermal insulation"],
    "Energy Consumption": ["energy-efficient equipment", "fuel-efficient vehicles", "energy optimization", "low-energy site operations", "reduced generator use", "hybrid construction machinery", "site energy plan"],
    "Thermal Performance": ["thermal envelope", "insulation", "U-value", "heat loss"],
    "Fuel Type for Equipment": ["Biodiesel", "alternative fuel", "low sulfur diesel", "renewable diesel", "clean fuel specification", "fuel switching", "emissions-compliant equipment", "non-fossil fuel use", "fuel quality standards"],
    "Lifecycle Carbon Reporting": ["EPD", "ISO 14040", "LCA", "carbon disclosure", "cradle to grave", "cradle to grave analysis", "life cycle assessment", "embodied carbon",  "global warming potential", "whole life carbon", "whole of life emissions"],
    "Low Emission Construction Materials": ["EPD certified", "climate-friendly materials", "green concrete", "green steel", "low GWP products", "low embodied carbon", "low emission materials", "low-carbon concrete", "recycled content", "recycled steel", "sustainable aggregates"],
    "Renewable Energy Systems": ["solar PV", "solar thermal", "on-site renewables", "wind turbine", "clean energy supply"],
    "Renewable Energy Use": ["solar PV", "wind turbine", "renewable sources", "clean energy"],
    "Scope 1 GHG Emissions - Onsite Emissions Reduction Measures": ["low-emission equipment", "electric construction machinery", "no-idling policy", "diesel alternatives"],
    "Scope 2 GHG Emissions - Procurement of Renewable or Clean Electricity": ["renewable electricity", "grid decarbonization", "clean energy supplier", "green power purchase"],
    "Waste Management in Construction": ["construction waste plan", "waste diversion", "recycling targets", "deconstruction waste", "waste audit", "material reuse"],
    "Ecological Impacts": ["biodiversity management plan", "ecological preservation", "flora and fauna protection", "habitat conservation", "ecological corridors", "species impact assessment", "no net loss of biodiversity", "critical habitat avoidance"],
    "Land Use Change": ["controlled site clearance", "habitat protection", "reduced land disturbance", "preservation of existing vegetation", "grading minimization", "sensitive site planning", "ecological buffer zones"],
    "Sustainable Maintenance Planning": ["maintenance plan", "O&M manual", "sustainable operations", "long-term performance", "building tuning"],
    "Air Quality (PM)": ["dust suppression", "PM10 control", "particulate mitigation", "air quality management plan", "water spraying", "dust barriers", "low-dust equipment", "site air monitoring", "fine particle control"],
    "Biological Oxygen Demand (BOD)": ["biological oxygen demand", "BOD limits", "wastewater treatment", "treated discharge", "water effluent quality", "oxygen-demanding substances", "construction wastewater control", "water discharge permit", "EIA water standards"],
    "Chemical Oxygen Demand (COD)": ["chemical oxygen demand", "COD threshold", "treated effluent", "wastewater treatment", "organic load reduction", "water discharge monitoring", "pollutant load control", "construction site effluent standards", "COD testing protocol"],
    "Light Pollution": ["glare control", "shielded lighting", "cut-off luminaires", "dark-sky compliant", "timers or sensors", "reduced spill lighting", "low-impact exterior lighting", "night sky protection"],
    "Noise Pollution": ["noise monitoring", "noise control plan", "sound barriers", "decibel limits", "acoustic insulation", "quiet equipment", "low-noise machinery"],
    "Soil Contamination": ["soil remediation", "contamination prevention", "heavy metals testing", "hazardous waste containment", "soil quality monitoring", "clean soil management", "protective earthworks", "baseline soil assessment", "EIA soil standards"],
    "Suspended Solids": ["suspended solids control", "TSS limits", "sediment traps", "water filtration", "silt fencing", "particle settling tank", "turbidity control", "sedimentation basin", "construction runoff management"],
    "pH Level": ["pH monitoring", "acidity control", "alkalinity limits", "pH adjustment", "neutralization basin", "discharge pH standards", "pH compliant effluent", "pH testing protocol", "pH range compliance"],
    "Stormwater Management": ["stormwater", "runoff", "green infrastructure", "rainwater capture", "stormwater runoff", "permeable pavement", "rain garden", "swale", "detention basin"],
    "Water Harvesting and Efficiency": ["greywater system", "rainwater harvesting", "water recycling", "low-flow fixtures", "potable water reduction"],
    "Indoor Environmental Quality": ["IEQ", "acoustic comfort", "air changes per hour", "comfort metrics", "daylight factor", "daylighting", "indoor air quality", "low VOC", "thermal comfort", "ventilation", "ventilation rate"],
    "Stakeholder Transparency": ["stakeholder communication", "project transparency", "public disclosure", "open reporting", "stakeholder engagement strategy", "information sharing with communities", "project updates to stakeholders", "public access to project data", "transparency commitment clause"],
    "Training and Capacity Building": ["construction workforce training", "capacity building plan", "upskilling program", "technical training for laborers", "site-based skills development", "vocational training", "certified training requirement", "on-the-job training", "education for site workers"],
    "Community Co-Design": ["community engagement", "participatory planning", "stakeholder consultation", "co-design process", "local stakeholder input", "community design workshops", "inclusive planning sessions", "collaborative design", "engagement with affected communities"],
    "Community Engagement": ["co-design", "community feedback", "community input", "feedback sessions", "participatory planning", "public consultation", "public meetings", "stakeholder consultation", "stakeholder input"],
    "Local Employment": ["community employment", "regional workforce", "local hiring", "community-based labor", "regional workforce participation", "employment of local residents", "priority to local workers", "community employment target", "inclusion of local subcontractors", "local job creation"],
    "Gender Inclusion": ["women participation", "female workforce", "gender equity", "women in construction", "female labor participation", "gender-inclusive hiring", "women employment target", "gender-responsive workforce plan", "gender balance in project teams", "inclusion of women-owned subcontractors", "gender diversity reporting"],
    "Gender Responsive Design": ["gender-inclusive design", "safe design for women", "gender-sensitive infrastructure", "female-friendly facilities", "women’s access and safety", "gender-informed site layout", "inclusive public space", "stakeholder feedback on gender needs", "universal design for gender inclusion"],
    "Inclusive Design & Accessibility": ["universal design", "accessible building", "disability access", "barrier-free", "inclusive space"],
    "Worker Health & Safety": ["occupational health and safety", "HSE plan", "personal protective equipment", "PPE compliance", "site safety management", "injury prevention", "safety training", "hazard control", "safety monitoring protocol", "zero accident policy"],
    "Health & Well-being (Indoor Air, Lighting, Acoustic)": ["indoor air quality", "daylighting", "low VOC", "thermal comfort", "acoustic comfort", "ventilation rates"],
    "Cost of Ecosystem Rehabilitation": ["restoration costs", "ecological rehabilitation", "green recovery"],
    "Cost of Relocation": ["resettlement costs", "displacement compensation"],
    "Building Information Modelling (BIM) Use": ["BIM", "BIM brief", "BIM coordination", "BIM execution plan", "building information modelling"],
    "Local Content and Sourcing": ["local procurement", "economic uplift", "regional impact", "local content requirement", "regionally sourced materials", "local suppliers", "community-based sourcing", "preference for local vendors", "domestic procurement target", "locally manufactured inputs", "use of local subcontractors"],
    "Local Economic Benefits": ["local economic development", "support for community enterprises", "local job creation", "inclusive procurement", "regional economic impact", "engagement of local businesses", "SME participation", "community-based suppliers", "local value retention"],
    "Circular Construction Practices": ["design for disassembly", "modular construction", "component reuse", "material passport", "circular design"],
    "Structure Durability": ["design life", "structural longevity", "durable infrastructure", "resilience to degradation", "maintenance-free period", "long-life materials", "infrastructure lifespan", "extended service life", "low-maintenance design"],
    "Lifecycle Cost Analysis": ["lifecycle cost analysis", "LCCA", "whole life costing", "long-term cost evaluation", "cost-benefit analysis", "maintenance cost forecasting", "total cost of ownership", "value for money over lifecycle"]
}

# --- Streamlit App ---
def main():
    st.title("Sustainability Assessment Tool (Multi-Format with OCR & NLP)")

    uploaded_files = st.file_uploader(
        "Upload documents/contracts (PDF, DOCX, XLSX)", 
        type=['pdf', 'docx', 'xlsx'],
        accept_multiple_files=True
    )

    st.sidebar.header("Manage Keywords")
    st.sidebar.info(
    """
    This application analyzes documents for sustainability indicators based on keyword families.
    NLP is used to find semantically similar phrases.
    Ensure Tesseract OCR and Poppler are installed and configured for scanned PDF/image processing.
    """
    )
    keyword_data = display_keyword_management()

    all_analysis_results = []
    if uploaded_files:
        for uploaded_file in uploaded_files:
            raw_text, segments_info, filename = extract_text_and_segment_info_from_file(uploaded_file)
            # Debug: Show extracted text for DOCX if needed
            # if filename.endswith('.docx'):
            #    st.text_area(f"Extracted text from {filename}", raw_text, height=200)

            if raw_text and raw_text.strip():
                analysis_results = analyze_document_nlp(raw_text, segments_info, filename, keyword_data)
                all_analysis_results.append(analysis_results)
            else:
                st.error(f"Could not extract meaningful text from {uploaded_file.name}.")
        
        if all_analysis_results:
            display_overall_results(all_analysis_results)
            display_detailed_results_nlp(all_analysis_results)

def extract_text_and_segment_info_from_file(uploaded_file) -> Tuple[str, List[Dict[str, Any]], str]:
    filename = uploaded_file.name
    file_extension = os.path.splitext(filename)[1].lower()
    file_bytes = uploaded_file.getvalue()

    full_text_content = ""
    segments_data: List[Dict[str, Any]] = [] 
    current_total_offset = 0
    
    # For Poppler on Windows, if not in PATH:
    # poppler_path_win = r"C:\path\to\poppler-xx.xx.x\bin" # Example, adjust to your Poppler bin path
    # else: poppler_path_win = None

    try:
        st.info(f"Processing {filename} (type: {file_extension})...")
        if file_extension == '.pdf':
            text_from_pypdf2 = ""
            pypdf2_successful = False
            try:
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
                if pdf_reader.is_encrypted:
                    try:
                        pdf_reader.decrypt('')
                    except Exception as decrypt_err:
                        st.warning(f"Could not decrypt PDF '{filename}': {decrypt_err}. OCR will be attempted if it's image-based.")

                temp_full_text = ""
                temp_segments_data = []
                temp_offset = 0
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text() or ""
                    temp_full_text += page_text + "\n"
                    temp_offset += len(page_text) + 1
                    temp_segments_data.append({"id": i + 1, "type": "Page (Text)", "text_end_offset": temp_offset})
                
                if temp_full_text.strip():
                    full_text_content = temp_full_text
                    segments_data = temp_segments_data
                    current_total_offset = temp_offset
                    pypdf2_successful = True
                    st.success(f"Extracted text from PDF '{filename}' using PyPDF2.")
            except Exception as e_pypdf2:
                st.warning(f"PyPDF2 failed for '{filename}': {e_pypdf2}. Will attempt OCR.")

            if not pypdf2_successful or not full_text_content.strip():
                st.info(f"Attempting OCR for PDF '{filename}'...")
                full_text_content = "" 
                segments_data = []
                current_total_offset = 0
                try:
                    # images = convert_from_bytes(file_bytes, dpi=200, poppler_path=poppler_path_win) # Use this if poppler_path_win is set
                    images = convert_from_bytes(file_bytes, dpi=200) 
                    for i, image in enumerate(images):
                        page_text = pytesseract.image_to_string(image) or ""
                        full_text_content += page_text + "\n"
                        current_total_offset += len(page_text) + 1
                        segments_data.append({"id": i + 1, "type": "Page (OCR)", "text_end_offset": current_total_offset})
                    if full_text_content.strip():
                        st.success(f"Extracted text from PDF '{filename}' using OCR.")
                    else:
                        st.warning(f"OCR for PDF '{filename}' yielded no text.")
                except Exception as e_ocr:
                    st.error(f"OCR failed for PDF '{filename}': {e_ocr}. Check Tesseract/Poppler installation (see script comments) and file integrity.")
                    return "", [], filename

        elif file_extension == '.docx':
            doc = DocxDocument(BytesIO(file_bytes))
            
            para_segment_counter = 1
            for para in doc.paragraphs:
                para_text = para.text or ""
                if para_text.strip():
                    full_text_content += para_text + "\n"
                    current_total_offset += len(para_text) + 1
                    segments_data.append({"id": f"Para_{para_segment_counter}", "type": "Paragraph", "text_end_offset": current_total_offset})
                    para_segment_counter +=1
            
            # Extract text from tables
            for table_idx, table in enumerate(doc.tables):
                # Optional: Add a header for the table in the text, though it might disrupt sentence tokenization
                # table_text_header = f"\n--- DOCX Table {table_idx+1} Start ---\n"
                # full_text_content += table_text_header
                # current_total_offset += len(table_text_header)
                
                for row_idx, row in enumerate(table.rows):
                    row_texts = []
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text_content = ""
                        # Iterate through paragraphs in a cell to get all text
                        for paragraph in cell.paragraphs:
                            cell_text_content += paragraph.text + " " # Add space between paragraphs in a cell
                        
                        cell_text_content = cell_text_content.strip()
                        if cell_text_content:
                             row_texts.append(cell_text_content)
                    
                    if row_texts:
                        # Join cell texts for the current row, then add as a new line.
                        # This treats each row's content more like a continuous text block.
                        concatenated_row_text = " | ".join(row_texts) + "\n" # Using " | " as a visual separator for cell content
                        full_text_content += concatenated_row_text
                        current_total_offset += len(concatenated_row_text)
                        # Create a segment for each row that has content
                        segments_data.append({"id": f"Table{table_idx+1}_Row{row_idx+1}", "type": "Table_Row", "text_end_offset": current_total_offset})

                # Optional: Add a footer for the table in the text
                # table_text_footer = f"\n--- DOCX Table {table_idx+1} End ---\n"
                # full_text_content += table_text_footer
                # current_total_offset += len(table_text_footer)

            # Basic image OCR from DOCX (using doc.part.rels)
            img_ocred_segment_counter = 1
            try:
                for rel_id in doc.part.rels:
                    rel = doc.part.rels[rel_id]
                    if "image" in rel.target_ref:
                        image_part = rel.target_part
                        image_bytes_docx = image_part.blob
                        img_obj = Image.open(BytesIO(image_bytes_docx))
                        ocr_text_from_image = pytesseract.image_to_string(img_obj) or ""
                        if ocr_text_from_image.strip():
                            separator = f"\n--- OCRd Content from DocImage {img_ocred_segment_counter} ---\n"
                            full_text_content += separator + ocr_text_from_image + "\n"
                            current_total_offset += len(separator) + len(ocr_text_from_image) + 1
                            segments_data.append({"id": f"DocImage_{img_ocred_segment_counter}", "type": "Image (OCR)", "text_end_offset": current_total_offset})
                            img_ocred_segment_counter += 1
            except Exception as e_docx_img_ocr:
                st.warning(f"Could not OCR some images in DOCX '{filename}': {e_docx_img_ocr}. Ensure Tesseract is configured.")
            
            if full_text_content.strip():
                 st.success(f"Extracted text from DOCX '{filename}' (including paragraphs, tables, and attempted image OCR).")
            else:
                 st.warning(f"Extraction from DOCX '{filename}' yielded little or no text. Check document structure and content.")


        elif file_extension == '.xlsx':
            workbook = openpyxl.load_workbook(BytesIO(file_bytes))
            sheet_num = 1
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                # Optional: Add sheet name as a header in the text.
                # sheet_header = f"\n--- Excel Sheet Start: {sheet_name} ---\n"
                # full_text_content += sheet_header
                # current_total_offset += len(sheet_header)
                
                sheet_cell_texts = []
                for row_idx, row in enumerate(sheet.iter_rows()):
                    row_text_parts = []
                    for cell_idx, cell in enumerate(row):
                        if cell.value is not None:
                            # Forcing string conversion and stripping whitespace
                            cell_str_value = str(cell.value).strip()
                            if cell_str_value: # Only add non-empty cell values
                                row_text_parts.append(cell_str_value) # Keep it simple for NLP
                    if row_text_parts:
                        # Join cells in a row with a space, then add a newline for the row
                        sheet_cell_texts.append(" ".join(row_text_parts)) 
                
                if sheet_cell_texts:
                    sheet_text_content = "\n".join(sheet_cell_texts) + "\n"
                    full_text_content += sheet_text_content
                    current_total_offset += len(sheet_text_content)
                    segments_data.append({"id": sheet_name, "type": "Sheet", "text_end_offset": current_total_offset})
                sheet_num +=1
            if full_text_content.strip():
                st.success(f"Extracted text from XLSX '{filename}'.")
            else:
                st.warning(f"Extraction from XLSX '{filename}' yielded no text.")
        
        else:
            st.error(f"Unsupported file type: {file_extension} for file {filename}")
            return "", [], filename
        
        # Sort segments by their end offset to ensure correct location finding
        segments_data.sort(key=lambda s: s["text_end_offset"])
        return full_text_content, segments_data, filename

    except Exception as e:
        st.error(f"General error extracting text from {filename} ({file_extension}): {e}")
        import traceback
        st.error(traceback.format_exc()) # More detailed error for debugging
        return "", [], filename


def find_segment_for_sentence_offset(sentence_start_offset: int, segments_info: List[Dict[str, Any]]) -> Dict[str, Any]:
    for segment_info in segments_info:
        if sentence_start_offset < segment_info["text_end_offset"]:
            return segment_info
    return segments_info[-1] if segments_info else {"id": "Unknown", "type": "Segment", "text_end_offset": float('inf')}


def analyze_document_nlp(document_text: str, segments_info: List[Dict[str, Any]], filename: str, indicator_keywords_matrix: Dict[str, List[str]]) -> Dict:
    detected_indicators_nlp = {}
    
    document_text_lower = document_text.lower() 
    document_sentences_original_case = nltk.sent_tokenize(document_text) 
    
    if not document_sentences_original_case:
        st.warning(f"No sentences were tokenized from {filename}. Document might be empty or unparseable.")
        return {
            "num_indicators": 0,
            "dimension_coverage": {"Environmental": 0, "Social": 0, "Economic": 0},
            "matched_indicators_nlp": {},
            "ambition_level": "Not Assessed / Very Low", # Corrected default
            "extracted_from": filename
        }

    document_sentences_for_embedding = [s.lower() for s in document_sentences_original_case]

    st.write(f"Embedding sentences for {filename} (Total: {len(document_sentences_for_embedding)})...")
    if not document_sentences_for_embedding: # Double check after lowercasing
        st.warning(f"No sentences available for embedding in {filename}.")
        # Return structure consistent with no indicators found
        return {
            "num_indicators": 0,
            "dimension_coverage": {"Environmental": 0, "Social": 0, "Economic": 0},
            "matched_indicators_nlp": {},
            "ambition_level": "Not Assessed / Very Low",
            "extracted_from": filename
        }

    doc_sentence_embeddings = np.array(similarity_model.encode(document_sentences_for_embedding, convert_to_tensor=False, show_progress_bar=True))
    if doc_sentence_embeddings.size == 0: # Check if embeddings are empty
        st.warning(f"Sentence embedding failed or resulted in empty array for {filename}.")
        return {
            "num_indicators": 0,
            "dimension_coverage": {"Environmental": 0, "Social": 0, "Economic": 0},
            "matched_indicators_nlp": {},
            "ambition_level": "Not Assessed / Very Low",
            "extracted_from": filename
        }


    st.write(f"Analyzing {filename} for keyword similarity...")
    progress_bar = st.progress(0)
    total_indicators_to_process = len(indicator_keywords_matrix)

    for i_indicator, (indicator, keyword_family) in enumerate(indicator_keywords_matrix.items()):
        for keyword_from_matrix in keyword_family:
            keyword_from_matrix_lower = keyword_from_matrix.lower()
            keyword_from_matrix_embedding = np.array(similarity_model.encode(keyword_from_matrix_lower, convert_to_tensor=False))
            
            # Ensure embeddings are 2D
            current_doc_embeddings = doc_sentence_embeddings
            if current_doc_embeddings.ndim == 1:
                 current_doc_embeddings = current_doc_embeddings.reshape(1, -1)
            if keyword_from_matrix_embedding.ndim == 1:
                 keyword_from_matrix_embedding = keyword_from_matrix_embedding.reshape(1, -1)

            if current_doc_embeddings.shape[0] == 0: # Should not happen if checked above, but safeguard
                continue

            similarities = cosine_similarity(
                keyword_from_matrix_embedding, 
                current_doc_embeddings      
            )[0] 

            for i_sentence, score in enumerate(similarities):
                if score >= SIMILARITY_THRESHOLD:
                    if i_sentence < len(document_sentences_original_case): # Boundary check
                        matched_sentence_original_case = document_sentences_original_case[i_sentence]
                        
                        sentence_start_offset = -1
                        try: 
                            sentence_start_offset = document_text.index(matched_sentence_original_case)
                        except ValueError: 
                            try: 
                                sentence_start_offset = document_text_lower.find(matched_sentence_original_case.lower())
                            except ValueError:
                                pass 

                        location_str = "Location unknown"
                        if sentence_start_offset != -1 and segments_info: # Ensure segments_info is not empty
                           segment_info = find_segment_for_sentence_offset(sentence_start_offset, segments_info)
                           location_str = f"{segment_info.get('type','Segment')} {segment_info.get('id', 'N/A')}"
                        
                        if indicator not in detected_indicators_nlp:
                            detected_indicators_nlp[indicator] = []
                        
                        existing_match = False
                        for entry in detected_indicators_nlp[indicator]:
                            if entry['original_keyword'] == keyword_from_matrix and \
                               entry['similar_phrase_in_doc'] == matched_sentence_original_case:
                                existing_match = True
                                if score > entry['similarity_score']: 
                                    entry['similarity_score'] = float(score)
                                    entry['location'] = location_str
                                break
                        
                        if not existing_match:
                            detected_indicators_nlp[indicator].append({
                                "original_keyword": keyword_from_matrix, 
                                "similar_phrase_in_doc": matched_sentence_original_case, 
                                "similarity_score": float(score),
                                "location": location_str,
                                "filename": filename 
                            })
            
            if indicator in detected_indicators_nlp: # Sort after all keywords for an indicator are processed
                detected_indicators_nlp[indicator].sort(key=lambda x: x['similarity_score'], reverse=True)
        
        progress_bar.progress((i_indicator + 1) / total_indicators_to_process if total_indicators_to_process > 0 else 1)


    num_indicators = len(detected_indicators_nlp)
    dimension_coverage = get_dimension_coverage(detected_indicators_nlp)
    ambition_level = get_ambition_level(num_indicators, dimension_coverage)

    return {
        "num_indicators": num_indicators,
        "dimension_coverage": dimension_coverage,
        "matched_indicators_nlp": detected_indicators_nlp,
        "ambition_level": ambition_level,
        "extracted_from": filename
    }

# --- get_dimension_coverage and get_ambition_level remain the same ---
def get_dimension_coverage(detected_indicators: Dict[str, List]) -> Dict[str, int]:
    environmental_indicators = [
        "Building Energy Efficiency", "Energy Consumption", "Thermal Performance", 
        "Fuel Type for Equipment", "Lifecycle Carbon Reporting", "Low Emission Construction Materials", 
        "Renewable Energy Systems", "Renewable Energy Use", 
        "Scope 1 GHG Emissions - Onsite Emissions Reduction Measures", 
        "Scope 2 GHG Emissions - Procurement of Renewable or Clean Electricity", 
        "Waste Management in Construction", "Ecological Impacts", "Land Use Change", 
        "Sustainable Maintenance Planning", "Air Quality (PM)", "Biological Oxygen Demand (BOD)", 
        "Chemical Oxygen Demand (COD)", "Light Pollution", "Noise Pollution", "Soil Contamination", 
        "Suspended Solids", "pH Level", "Stormwater Management", "Water Harvesting and Efficiency", 
        "Indoor Environmental Quality"
    ]
    social_indicators = [
        "Stakeholder Transparency", "Training and Capacity Building", "Community Co-Design", 
        "Community Engagement", "Local Employment", "Gender Inclusion", "Gender Responsive Design", 
        "Inclusive Design & Accessibility", "Worker Health & Safety", 
        "Health & Well-being (Indoor Air, Lighting, Acoustic)" 
    ]
    economic_indicators = [
        "Cost of Ecosystem Rehabilitation", "Cost of Relocation", 
        "Building Information Modelling (BIM) Use", 
        "Local Content and Sourcing", "Local Economic Benefits", 
        "Circular Construction Practices", 
        "Structure Durability", 
        "Lifecycle Cost Analysis" 
    ]
    
    environmental_count = 0
    social_count = 0
    economic_count = 0

    detected_indicator_names = set(detected_indicators.keys())

    for ind_name in detected_indicator_names:
        if ind_name in environmental_indicators:
            environmental_count +=1
        elif ind_name in social_indicators: 
            social_count +=1
        elif ind_name in economic_indicators: 
            economic_count +=1
            
    return {
        "Environmental": environmental_count,
        "Social": social_count,
        "Economic": economic_count
    }

def get_ambition_level(num_indicators: int, dimension_coverage: Dict[str, int]) -> str:
    num_dimensions_spanned = 0
    if dimension_coverage.get("Environmental", 0) > 0:
        num_dimensions_spanned += 1
    if dimension_coverage.get("Social", 0) > 0:
        num_dimensions_spanned += 1
    if dimension_coverage.get("Economic", 0) > 0:
        num_dimensions_spanned += 1

    if num_indicators == 0:
        return "Not Assessed / Very Low"
    if num_indicators >= 10 and num_dimensions_spanned == 3:
        return "High"
    elif (num_indicators >= 10 and num_dimensions_spanned < 3) or \
         (5 <= num_indicators <= 9 and num_dimensions_spanned >= 2):
        return "Medium"
    elif 1 <= num_indicators <= 4: 
        return "Low"
    elif num_indicators > 0: 
        return "Low" 
    else: 
        return "Not Assessed / Very Low"

# --- display_overall_results remains largely the same ---
def display_overall_results(all_analysis_results: List[Dict]):
    st.header("Overall Analysis Summary")

    st.subheader("Overall Ambition Level Distribution")
    if all_analysis_results:
        # Filter out None or empty ambition levels if any document failed very early
        ambition_levels_present = sorted(list(set(res['ambition_level'] for res in all_analysis_results if res.get('ambition_level'))))
        
        if not ambition_levels_present:
            st.write("No ambition data to display (all documents may have failed processing before ambition scoring).")
            return


        overall_ambition_data = pd.DataFrame({
            'Level': ambition_levels_present,
            'Count': [sum(1 for res in all_analysis_results if res.get('ambition_level') == level) for level in ambition_levels_present]
        })
        if not overall_ambition_data.empty and 'Count' in overall_ambition_data and overall_ambition_data['Count'].sum() > 0 :
            color_scale = alt.Scale(
                domain=['Not Assessed / Very Low', 'Low', 'Medium', 'High'],
                range=['#d3d3d3', '#ffcc00', '#ff6600', '#33cc33'] 
            )
            pie_chart = alt.Chart(overall_ambition_data).mark_arc(outerRadius=120).encode(
                theta=alt.Theta(field="Count", type="quantitative"),
                color=alt.Color(field="Level", type="nominal", scale=color_scale, sort=ambition_levels_present),
                tooltip=['Level', 'Count']
            ).properties(title="Ambition Level Distribution")
            st.altair_chart(pie_chart, use_container_width=True)
        else:
            st.write("No ambition data to display or all counts are zero.")

        if len(all_analysis_results) > 0:
            # Ensure num_indicators exists before summing
            valid_results_for_avg = [res['num_indicators'] for res in all_analysis_results if 'num_indicators' in res]
            if valid_results_for_avg:
                 avg_indicators = sum(valid_results_for_avg) / len(valid_results_for_avg)
                 st.metric(label="Average Number of Indicators Detected per Document", value=f"{avg_indicators:.2f}")
            else:
                st.metric(label="Average Number of Indicators Detected per Document", value="N/A (No valid results)")
        else:
            st.metric(label="Average Number of Indicators Detected per Document", value="N/A")


        st.subheader("Overall Dimension Coverage")
        overall_dimension_data = {
            'Environmental': sum(res.get('dimension_coverage', {}).get('Environmental', 0) for res in all_analysis_results),
            'Social': sum(res.get('dimension_coverage', {}).get('Social', 0) for res in all_analysis_results),
            'Economic': sum(res.get('dimension_coverage', {}).get('Economic', 0) for res in all_analysis_results)
        }
        dimension_df = pd.DataFrame(list(overall_dimension_data.items()), columns=['Dimension', 'Count'])
        if not dimension_df.empty and dimension_df['Count'].sum() > 0:
            chart = alt.Chart(dimension_df).mark_bar().encode(
                x=alt.X('Dimension', sort=None),
                y='Count',
                color=alt.Color('Dimension', legend=None),
                tooltip=['Dimension', 'Count']
            ).properties(title="Combined Dimension Coverage (Total Indicators Found)")
            st.altair_chart(chart, use_container_width=True)
        else:
            st.write("No dimension coverage data to display.")
    else:
        st.write("No analysis results to display overall summary.")


def display_detailed_results_nlp(all_analysis_results: List[Dict]):
    st.header("Detailed Results per Document (NLP Similarity)")

    for result in all_analysis_results:
        # Ensure essential keys exist before trying to display
        if not all(k in result for k in ['extracted_from', 'ambition_level', 'num_indicators', 'dimension_coverage', 'matched_indicators_nlp']):
            st.warning(f"Skipping display for a document due to missing analysis data (possibly {result.get('extracted_from', 'Unknown File')}).")
            continue

        with st.expander(f"Analysis of: {result['extracted_from']}", expanded=False): 
            col1, col2 = st.columns(2)
            with col1:
                st.metric(label="Ambition Level", value=result['ambition_level'])
            with col2:
                st.metric(label="Number of Unique Indicators Detected", value=result['num_indicators'])

            st.subheader("Dimension Coverage for this Document")
            dimension_data = pd.DataFrame(list(result["dimension_coverage"].items()), columns=['Dimension', 'Count'])
            if not dimension_data.empty and dimension_data['Count'].sum() > 0 :
                chart = alt.Chart(dimension_data).mark_bar().encode(
                    x=alt.X('Dimension', sort=None),
                    y='Count',
                    color=alt.Color('Dimension', legend=None),
                    tooltip=['Dimension', 'Count']
                ).properties(height=300)
                st.altair_chart(chart, use_container_width=True)
            else:
                st.write("No dimension coverage to display for this document (no indicators found).")

            st.subheader(f"Matched Indicators and Similar Phrases (NLP) with {SIMILARITY_THRESHOLD*100:.0f}% threshold")
            if result["matched_indicators_nlp"]:
                data_nlp = []
                for indicator, matches_list in result["matched_indicators_nlp"].items():
                    for match_item in matches_list:
                        data_nlp.append({
                            "Indicator": indicator,
                            "Original Keyword (from Matrix)": match_item["original_keyword"],
                            "Detected Similar Phrase (from Contract)": match_item["similar_phrase_in_doc"],
                            "Similarity (%)": f"{match_item['similarity_score']*100:.2f}%",
                            "Location (Segment ID)": match_item["location"], 
                        })
                if data_nlp:
                    df_nlp = pd.DataFrame(data_nlp)
                    st.dataframe(df_nlp)
                else: 
                    st.write("No similar phrases found meeting the threshold for this document.")
            else:
                st.write("No indicators matched using NLP for this document.")
        st.markdown("---")


def display_keyword_management():
    st.sidebar.subheader("Indicators and Keyword Families") 
    
    if 'editable_keyword_data' not in st.session_state:
        st.session_state.editable_keyword_data = initial_indicator_keywords.copy()

    with st.sidebar.container(height=400): 
        for indicator, keywords in list(st.session_state.editable_keyword_data.items()): # Use list() for safe iteration if modifying dict
            with st.expander(indicator):
                keyword_string = ", ".join(keywords)
                new_keyword_str = st.text_area(
                    f"Keywords for {indicator}:", 
                    value=keyword_string, 
                    key=f"text_area_kw_{indicator.replace(' ', '_').replace('&', 'and').replace('(', '').replace(')', '').replace('.', '')}" # Sanitize key
                )
                if new_keyword_str != keyword_string:
                    st.session_state.editable_keyword_data[indicator] = [k.strip() for k in new_keyword_str.split(",") if k.strip()]
                
                if st.button(f"Remove: {indicator}", key=f"remove_btn_{indicator.replace(' ', '_').replace('&', 'and').replace('(', '').replace(')', '').replace('.', '')}"): # Sanitize key
                    if indicator in st.session_state.editable_keyword_data:
                        del st.session_state.editable_keyword_data[indicator]
                        st.rerun()

    st.sidebar.subheader("Add New Indicator")
    new_indicator_name = st.sidebar.text_input("New Indicator Name", key="new_indicator_name_input_sidebar")
    new_indicator_kws = st.sidebar.text_area("Keywords (comma-separated)", key="new_indicator_kws_input_sidebar")
    if st.sidebar.button("Add Indicator", key="add_new_indicator_button_sidebar"):
        if new_indicator_name and new_indicator_kws:
            if new_indicator_name not in st.session_state.editable_keyword_data:
                st.session_state.editable_keyword_data[new_indicator_name] = [k.strip() for k in new_indicator_kws.split(",") if k.strip()]
                st.sidebar.success(f"Added: {new_indicator_name}")
                # Consider clearing inputs here:
                # st.session_state.new_indicator_name_input_sidebar = "" 
                # st.session_state.new_indicator_kws_input_sidebar = ""
                st.rerun()
            else:
                st.sidebar.error("Indicator already exists.")
        else:
            st.sidebar.error("Indicator name and keywords are required.")
            
    return st.session_state.editable_keyword_data


if __name__ == "__main__":
    main()